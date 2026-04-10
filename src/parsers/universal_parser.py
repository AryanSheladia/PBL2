from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import fitz  # PyMuPDF
import docx
import pandas as pd
from rapidfuzz import fuzz

from src.models.schema import ParsedDocument, Section
from src.core.doc_type import detect_doc_type
from src.core.fingerprint import fingerprint


# -------------------------------------------------------------------
# Templates
# -------------------------------------------------------------------

TEMPLATE_MAP = {
    "PRD": "src/templates/pubmatic/template_prd.json",
    "TDD": "src/templates/pubmatic/template_tdd_adr.json",
    "API": "src/templates/pubmatic/template_api_spec.json",
    "RUNBOOK": "src/templates/pubmatic/template_runbook.json",
    "DATA": "src/templates/pubmatic/template_data_doc.json",
    "GENERIC": "src/templates/pubmatic/template_generic.json",  # REQUIRED
}


def load_template(path: str | Path) -> Dict[str, Any]:
    return json.loads(Path(path).read_text(encoding="utf-8"))


def _load_all_templates() -> List[Dict[str, Any]]:
    return [load_template(p) for p in TEMPLATE_MAP.values()]


def _merge_templates(templates: List[Dict[str, Any]]) -> Dict[str, Any]:
    """
    Merge N templates.
    Union by section_id, merge aliases.
    """
    merged = {
        "template_name": "+".join(t.get("template_name", "t") for t in templates),
        "doc_type": "+".join(t.get("doc_type", "") for t in templates if t.get("doc_type")),
        "sections": [],
    }

    by_id: Dict[str, set] = {}
    for t in templates:
        for s in t.get("sections", []):
            sid = s["section_id"]
            by_id.setdefault(sid, set()).update(s.get("aliases", []))

    for sid, aliases in by_id.items():
        merged["sections"].append({"section_id": sid, "aliases": sorted(list(aliases))})

    return merged


# -------------------------------------------------------------------
# Text utils + heading detection
# -------------------------------------------------------------------

def _norm(s: str) -> str:
    return re.sub(r"\s+", " ", s.strip().lower())


def _heading_like(line: str) -> bool:
    line = line.strip()
    if not line or len(line) > 140:
        return False
    if line.endswith("."):
        return False
    if len(line.split()) > 14:
        return False

    if re.match(r"^(\d+(\.\d+)*)\s+.+", line):
        return True

    letters = [c for c in line if c.isalpha()]
    if letters:
        caps_ratio = sum(1 for c in letters if c.isupper()) / len(letters)
        if caps_ratio >= 0.65:
            return True

    words = [w for w in line.split() if any(ch.isalpha() for ch in w)]
    if words:
        starts_caps = sum(1 for w in words if w[:1].isupper())
        if starts_caps / len(words) >= 0.6:
            return True

    return False


def _match_section_id(heading: str, template: Dict[str, Any]) -> Tuple[str, float]:
    h = _norm(heading)
    best_id = "UNMAPPED"
    best_score = 0

    for sec in template.get("sections", []):
        sid = sec["section_id"]
        for a in sec.get("aliases", []):
            score = fuzz.partial_ratio(h, _norm(a))
            if score > best_score:
                best_score = score
                best_id = sid

    conf = best_score / 100.0
    if best_score < 75:
        return "UNMAPPED", conf
    return best_id, conf


def _stable_section_id(file_path: Path, heading: str, existing_ids: set) -> str:
    slug = re.sub(r"[^a-z0-9]+", "-", _norm(heading)).strip("-")

    if not slug:
        slug = "unknown-section"

    base_id = f"{file_path.stem}::{slug}"
    section_id = base_id

    counter = 1
    while section_id in existing_ids:
        section_id = f"{base_id}-{counter}"
        counter += 1

    existing_ids.add(section_id)
    return section_id

def _anchor_from_heading(section_id: str, heading: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "-", _norm(heading)).strip("-")
    if not slug:
        slug = section_id.lower()
    return f"{section_id.lower()}::{slug}"


# -------------------------------------------------------------------
# Template scoring + multi-template selection (Enterprise)
# -------------------------------------------------------------------

def _template_score(template: Dict[str, Any], headings: List[str], full_text: str) -> float:
    score = 0.0
    ft = _norm(full_text)

    # Heading matches are high value
    for h in headings:
        sid, conf = _match_section_id(h, template)
        if sid != "UNMAPPED":
            score += 2.0 * conf

    # Text keyword hits are low value
    for sec in template.get("sections", []):
        for a in sec.get("aliases", []):
            if _norm(a) in ft:
                score += 0.25

    return score


def _choose_enterprise_template(full_text: str, candidates: List[Tuple[int, str, Optional[int]]]) -> Tuple[str, Dict[str, Any], Dict[str, float]]:
    """
    Choose best template by scoring all, then:
    - merge top K (K<=3) if their scores are close
    - ALWAYS include GENERIC as a backstop for leftover headings
    Returns: (doc_type_guess, merged_template, score_debug)
    """
    doc_type_guess = detect_doc_type(full_text)
    headings = [h for (_, h, _) in candidates] if candidates else []

    templates = _load_all_templates()

    scored = []
    for t in templates:
        s = _template_score(t, headings, full_text)
        scored.append((t, s))

    scored.sort(key=lambda x: x[1], reverse=True)

    # Always keep GENERIC around as a backstop template
    generic = next((t for t in templates if t.get("doc_type") == "GENERIC" or "generic" in t.get("template_name", "").lower()), None)

    best, best_score = scored[0]
    selected = [best]

    # Merge up to 3 if close enough
    for (t, sc) in scored[1:]:
        if len(selected) >= 3:
            break
        if best_score > 0 and sc >= 0.80 * best_score and sc > 0:
            selected.append(t)

    if generic and generic not in selected:
        selected.append(generic)

    merged = _merge_templates(selected)

    debug_scores = {t.get("template_name", "t"): sc for (t, sc) in scored[:5]}
    return doc_type_guess, merged, debug_scores


# -------------------------------------------------------------------
# PDF extraction (font-aware headings)
# -------------------------------------------------------------------

def _extract_pdf(pdf_path: Path) -> Tuple[str, List[Dict[str, Any]], List[Tuple[int, int, int]]]:
    doc = fitz.open(str(pdf_path))
    all_lines: List[Dict[str, Any]] = []
    parts: List[str] = []
    page_map: List[Tuple[int, int, int]] = []
    cursor = 0

    for page_no, page in enumerate(doc, start=1):
        d = page.get_text("dict")
        page_lines = []

        for block in d.get("blocks", []):
            for line in block.get("lines", []):
                spans = line.get("spans", [])
                if not spans:
                    continue
                text = "".join(s.get("text", "") for s in spans).strip()
                if not text:
                    continue

                max_size = max(float(s.get("size", 0)) for s in spans)
                bold = any("bold" in str(s.get("font", "")).lower() for s in spans)

                all_lines.append({"text": text, "page": page_no, "size": max_size, "bold": bold})
                page_lines.append(text)

        page_text = "\n".join(page_lines) + "\n"
        start = cursor
        parts.append(page_text)
        cursor += len(page_text)
        end = cursor
        page_map.append((start, end, page_no))

    return "".join(parts), all_lines, page_map


def _pdf_heading_candidates(lines: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    if not lines:
        return []
    sizes = sorted([l["size"] for l in lines])
    body = sizes[len(sizes) // 2]
    out = []

    for l in lines:
        if not _heading_like(l["text"]):
            continue
        if l["size"] >= body + 2 or (l["bold"] and l["size"] >= body + 0.5):
            out.append(l)

    return out


def _build_candidates_pdf(full_text: str, lines: List[Dict[str, Any]]) -> List[Tuple[int, str, Optional[int]]]:
    candidates: List[Tuple[int, str, Optional[int]]] = []
    hc = _pdf_heading_candidates(lines)
    for h in hc:
        idx = full_text.find(h["text"])
        if idx != -1:
            candidates.append((idx, h["text"], h["page"]))
    return sorted(set(candidates), key=lambda x: x[0])


def _char_range_to_pages(page_map: Optional[List[Tuple[int, int, int]]], start: int, end: int) -> Optional[Tuple[int, int]]:
    if not page_map:
        return None
    pages = set()
    for s, e, p in page_map:
        if e <= start:
            continue
        if s >= end:
            break
        pages.add(p)
    return (min(pages), max(pages)) if pages else None


# -------------------------------------------------------------------
# DOCX extraction (structure-aware headings)
# -------------------------------------------------------------------

def _extract_docx_with_structure(docx_path: Path) -> Tuple[str, List[Dict[str, Any]]]:
    d = docx.Document(str(docx_path))
    lines: List[Dict[str, Any]] = []
    full_parts: List[str] = []

    for p in d.paragraphs:
        text = p.text.strip()
        if not text:
            continue

        style_name = p.style.name if p.style else ""
        is_heading_style = "heading" in style_name.lower()
        bold_run = any((run.text and run.bold) for run in p.runs)

        is_short = len(text.split()) <= 12
        numbered = bool(re.match(r"^(\d+(\.\d+)*)\s+.+", text))

        is_heading = is_heading_style or (bold_run and is_short) or numbered or _heading_like(text)

        lines.append({"text": text, "is_heading": is_heading})
        full_parts.append(text + "\n")

    return "".join(full_parts), lines


def _build_candidates_docx(full_text: str, docx_lines: List[Dict[str, Any]]) -> List[Tuple[int, str, Optional[int]]]:
    candidates: List[Tuple[int, str, Optional[int]]] = []
    cursor = 0
    for item in docx_lines:
        text = item["text"]
        idx = full_text.find(text, cursor)
        if idx != -1 and item["is_heading"]:
            candidates.append((idx, text, None))
            cursor = idx + len(text)
        elif idx != -1:
            cursor = idx + len(text)

    return sorted(set(candidates), key=lambda x: x[0])


# -------------------------------------------------------------------
# TXT / CSV extraction
# -------------------------------------------------------------------

def _extract_txt(txt_path: Path) -> str:
    return txt_path.read_text(encoding="utf-8", errors="ignore")


def _extract_csv(csv_path: Path) -> str:
    df = pd.read_csv(csv_path)
    lines = []
    lines.append("DATASET OVERVIEW")
    lines.append(f"Rows: {len(df)}")
    lines.append(f"Columns ({len(df.columns)}): {', '.join(map(str, df.columns[:30]))}" + ("..." if len(df.columns) > 30 else ""))
    lines.append("\nSCHEMA\n" + "\n".join([f"- {c}" for c in df.columns[:50]]))
    lines.append("\nSAMPLE QUERIES\nSELECT * FROM dataset LIMIT 5;")
    lines.append("\nSAMPLE\n" + df.head(5).to_string(index=False))
    return "\n".join(lines) + "\n"


# -------------------------------------------------------------------
# Semantic fallback
# -------------------------------------------------------------------

def _split_into_chunks(full_text: str) -> List[str]:
    chunks = [p.strip() for p in re.split(r"\n\s*\n", full_text) if p.strip()]
    if len(chunks) >= 3:
        return chunks

    chunks = []
    buff: List[str] = []
    for line in full_text.splitlines():
        line = line.strip()
        if not line:
            continue

        if re.match(r"^(\-|\*|•|\d+\.|\d+\))\s+", line) and buff:
            chunks.append(" ".join(buff).strip())
            buff = [line]
        else:
            buff.append(line)

    if buff:
        chunks.append(" ".join(buff).strip())

    return [c for c in chunks if c]


def _semantic_bucket(full_text: str, template: Dict[str, Any], file_path: Path) -> List[Section]:
    chunks = _split_into_chunks(full_text)

    sec_ids = [s["section_id"] for s in template.get("sections", [])]
    misc_id = "MISC_NOTES" if "MISC_NOTES" in sec_ids else (sec_ids[-1] if sec_ids else "MISC_NOTES")

    buckets: Dict[str, List[str]] = {sid: [] for sid in sec_ids}
    buckets.setdefault(misc_id, [])

    def score(chunk: str, sec: Dict[str, Any]) -> int:
        ct = _norm(chunk)
        s = 0
        for a in sec.get("aliases", []):
            if _norm(a) in ct:
                s += 1
        return s

    for chunk in chunks:
        best_sid = misc_id
        best = 0
        for sec in template.get("sections", []):
            sid = sec["section_id"]
            sc = score(chunk, sec)
            if sc > best:
                best = sc
                best_sid = sid
        if best == 0:
            best_sid = misc_id
        buckets[best_sid].append(chunk)

    sections: List[Section] = []
    existing_ids = set()
    for i, sid in enumerate(sec_ids):
        if not buckets.get(sid):
            continue
        content = "\n\n".join(buckets[sid]).strip()
        sections.append(
            Section(
                section_id=_stable_section_id(file_path, sid, existing_ids),
                heading=sid.replace("_", " ").title(),
                anchor=f"{sid.lower()}::semantic",
                content=content,
                char_range=(0, 0),
                confidence=0.55,
                fingerprint=fingerprint(content),
                meta={"mode": "semantic_fallback"},
            )
        )

    # always output misc if it has content and isn’t already emitted
    if misc_id not in sec_ids and buckets.get(misc_id):
        content = "\n\n".join(buckets[misc_id]).strip()
        sections.append(
            Section(
                section_id=misc_id,
                heading=misc_id.replace("_", " ").title(),
                anchor=f"{misc_id.lower()}::semantic",
                content=content,
                char_range=(0, 0),
                confidence=0.50,
                fingerprint=fingerprint(content),
                meta={"mode": "semantic_fallback"},
            )
        )

    return sections


# -------------------------------------------------------------------
# Main
# -------------------------------------------------------------------

def parse_any(file_path: str | Path) -> ParsedDocument:
    file_path = Path(file_path)
    suffix = file_path.suffix.lower()

    full_text: str
    page_map: Optional[List[Tuple[int, int, int]]] = None
    candidates: List[Tuple[int, str, Optional[int]]] = []

    if suffix == ".pdf":
        full_text, lines, page_map = _extract_pdf(file_path)
        candidates = _build_candidates_pdf(full_text, lines)

    elif suffix == ".docx":
        full_text, docx_lines = _extract_docx_with_structure(file_path)
        page_map = None
        candidates = _build_candidates_docx(full_text, docx_lines)

    elif suffix in [".txt", ".md"]:
        full_text = _extract_txt(file_path)
        page_map = None
        candidates = []

    elif suffix == ".csv":
        full_text = _extract_csv(file_path)
        page_map = None
        candidates = []

    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    # Enterprise template selection (score + merge + generic backstop)
    doc_type_guess, template, _debug_scores = _choose_enterprise_template(full_text, candidates)

    # If no headings -> semantic fallback
    if len(candidates) == 0:
        sections = _semantic_bucket(full_text, template, file_path)
        return ParsedDocument(
            source_path=str(file_path),
            template_name=template.get("template_name", "unknown_template"),
            doc_type=doc_type_guess,
            full_text=full_text,
            sections=sections,
        )

    # Heading split (and force UNMAPPED into GENERIC/MISC_NOTES)
    sections: List[Section] = []
    existing_ids = set()
    for i, (start, heading, page) in enumerate(candidates):
        end = candidates[i + 1][0] if i + 1 < len(candidates) else len(full_text)
        chunk = full_text[start:end].strip()

        sid, conf = _match_section_id(heading, template)

        # Enterprise: never leave UNMAPPED; fallback to MISC_NOTES
        if sid == "UNMAPPED":
            sid = "MISC_NOTES"
            conf = max(conf, 0.55)

        anchor = _anchor_from_heading(sid, heading)
        pr = _char_range_to_pages(page_map, start, end) if page_map else None

        sections.append(
            Section(
                section_id=_stable_section_id(file_path, heading, existing_ids),
                heading=heading,
                anchor=anchor,
                content=chunk,
                char_range=(start, end),
                page_range=pr,
                confidence=max(0.60, conf),
                fingerprint=fingerprint(chunk),
                meta={"mode": "heading_split", "doc_type_guess": doc_type_guess},
            )
        )

    return ParsedDocument(
        source_path=str(file_path),
        template_name=template.get("template_name", "unknown_template"),
        doc_type=doc_type_guess,
        full_text=full_text,
        sections=sections,
    )